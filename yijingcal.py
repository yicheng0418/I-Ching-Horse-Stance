from docx import Document
import random
import datetime
from detailhexagram import second_main
from lunarDateBaziCal import gan, zhi, lunar_months, lunar_days, gan_wuxing, zhi_wuxing
from lunarDateBaziCal import get_zodiac, get_ganzhi, get_day_ganzhi, get_hour_ganzhi

# 定义六个爻位的阴阳符号
YIN = "阴"
YANG = "阳"

BAGUA_YAO_MAPPING = {
    "乾": ["阳", "阳", "阳"],
    "兑": ["阴", "阳", "阳"],
    "离": ["阳", "阴", "阳"],
    "震": ["阴", "阴", "阳"],
    "巽": ["阳", "阳", "阴"],
    "坎": ["阴", "阳", "阴"],
    "艮": ["阳", "阴", "阴"],
    "坤": ["阴", "阴", "阴"]
}

# 定义先天八卦与数字的对应关系
BAGUA_MAPPING = {
    1: "乾",
    2: "兑",
    3: "离",
    4: "震",
    5: "巽",
    6: "坎",
    7: "艮",
    8: "坤"
}

# 定义后天八卦与数字的对应关系
HOUTIAN_BAGUA_MAPPING = {
    1: "坎",
    2: "坤",
    3: "震",
    4: "巽",
    5: "中宫",
    6: "乾",
    7: "兑",
    8: "艮",
    9: "离"
}

# 读取docx文件中的卦爻信息
def read_gua_yao_from_docx(file_path):
    doc = Document(file_path)
    gua_yao_mapping = {}
    table = doc.tables[0]  # 假设卦爻信息在第一个表格中
    for row in table.rows[1:]:  # 跳过表头行，从第二行开始读取数据
        cells = row.cells
        if len(cells) > 1:
            gua_name = cells[0].text.strip()
            gua_yao = [cell.text.strip() for cell in cells[1:]]  # 修改为列表形式
            gua_yao_mapping[gua_name] = gua_yao
    return gua_yao_mapping

# 根据阴阳爻位获取卦名
def get_gua_name_by_yao(gua_yao_mapping, hexagram):
    """根据爻位列表获取六十四卦的卦名"""
    for gua_name, gua_yao in gua_yao_mapping.items():
        if gua_yao == hexagram:
            return gua_name
    return None

# 修改后的生成本卦函数（使用新的逻辑）
def generate_base_hexagram_new():
    upper_num = random.randint(100, 999)
    lower_num = random.randint(100, 999)

    upper_hexagram_num = upper_num % 8 or 8
    lower_hexagram_num = lower_num % 8 or 8

    upper_gua = BAGUA_MAPPING[upper_hexagram_num]
    lower_gua = BAGUA_MAPPING[lower_hexagram_num]

    # 组合上卦和下卦，确保爻位的顺序（第1爻在底部，第6爻在顶部）
    return BAGUA_YAO_MAPPING[lower_gua] + BAGUA_YAO_MAPPING[upper_gua]

# 修改后的生成变爻函数（使用新的逻辑）
def generate_changing_line_new(base_hexagram):
    changing_num = random.randint(100, 999)
    changing_line_index = changing_num % 6 

    # 确定变爻的阴阳性质（考虑易经的爻位顺序）
    changing_line = YIN if base_hexagram[5 - changing_line_index] == YANG else YANG
    return changing_line, changing_line_index

# 根据变爻位生成变卦的函数
def generate_changing_hexagram(base_hexagram, changing_line, changing_line_index):
    changing_hexagram = base_hexagram.copy()
    # 考虑易经的爻位顺序，第1爻在底部，第6爻在顶部
    # 变爻位的索引需要调整为 5 - changing_line_index
    adjusted_index = 5 - changing_line_index
    # 将变爻位的阴阳性质应用到本卦中的相应爻位
    changing_hexagram[adjusted_index] = changing_line
    return changing_hexagram

# Assuming the lunarDateBaziCal provides functions for the lunar date and month
def generate_daily_hexagram_lunar(docx_path, gua_yao_mapping):
    # Get today's lunar date (you can use your lunarDateBaziCal logic here)
    today = datetime.datetime.now()
    
    # Get the lunar date (e.g., 20th day of the 9th month)
    lunar_day = lunar_days[today.day % 4]  # Lunar day (could be "初", "十", "廿", "卅")
    lunar_month = lunar_months[(today.month - 1) % 12]  # Lunar month (e.g., "正", "二", "三")

    # Combine lunar day and month as a seed
    lunar_seed = today.year * 100 + today.month * 10 + today.day  # You can modify the seed logic
    random.seed(lunar_seed)  # Use lunar date as the seed for randomization

    # Generate base hexagram using lunar seed
    base_hexagram = generate_base_hexagram_new()  # Use existing logic
    base_gua_name = get_gua_name_by_yao(gua_yao_mapping, base_hexagram)
    base_gua_content = extract_guaci_content(docx_path, base_gua_name)

    # Generate changing line and changing hexagram
    changing_line, changing_line_index = generate_changing_line_new(base_hexagram)
    changing_hexagram = generate_changing_hexagram(base_hexagram, changing_line, changing_line_index)
    changing_gua_name = get_gua_name_by_yao(gua_yao_mapping, changing_hexagram)
    changing_gua_content = extract_guaci_content(docx_path, changing_gua_name)

    # Return the result
    return {
        "base_gua": base_gua_name,
        "base_gua_content": base_gua_content,
        "changing_gua": changing_gua_name,
        "changing_gua_content": changing_gua_content
    }

# 提取卦辞内容
def extract_guaci_content(docx_path, gua_name):
    doc = Document(docx_path)
    guaci_content = ""

    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if gua_name in text:
            guaci_content = text
            break

    return guaci_content

# 根据输入的数字计算变爻位置
def calculate_changing_line(input_values):
    """根据输入值计算变爻位置"""
    total_sum = sum(input_values)
    moving_line = total_sum % 6
    if moving_line == 0:
        moving_line = 6
    return moving_line

# 计算变卦：应用变爻并更新卦象
def apply_changing_line_to_hexagram(hexagram, changing_line):
    hexagram_list = hexagram.copy()
    changing_line_index = 6 - changing_line  # 转换为0基索引
    hexagram_list[changing_line_index] = "阳" if hexagram_list[changing_line_index] == "阴" else "阴"
    return hexagram_list

# 根据输入的三个数字获取六爻卦信息
def number_to_hexagram_from_user_input(docx_path, gua_yao_mapping):
    user_input = input("请输入三个数字，以空格分隔，用于六爻数字占卜：")
    
    # 判断用户输入是否包含空格
    if ' ' in user_input:
        nums = list(map(int, user_input.split()))
    else:
        if len(user_input) != 3:
            print("请输入三个数字！")
            return
        nums = list(map(int, user_input))  # 将字符串拆分成单个字符并转化为数字
    
    if len(nums) != 3:
        print("请确保输入三个数字！")
        return
    
    # 获取上下卦的数字，确保范围在1-8之间
    lower = nums[0] % 8  # 计算下卦，确保其在 1-8 的范围内
    upper = (nums[1] + nums[2]) % 8  # 计算上卦，第二和第三个数字的和

    # 如果 upper 和 lower 结果为 0，需要转换为 8
    if lower == 0:
        lower = 8
    if upper == 0:
        upper = 8

    # 合并上下卦的爻位
    lower_gua_name = BAGUA_MAPPING.get(lower)
    upper_gua_name = BAGUA_MAPPING.get(upper)

    if not lower_gua_name or not upper_gua_name:
        print("无效的卦名！")
        return
    
    hexagram = BAGUA_YAO_MAPPING.get(upper_gua_name) + BAGUA_YAO_MAPPING.get(lower_gua_name)  # 返回六个爻的列表
    
    if not hexagram:
        print("无法合并上卦和下卦的爻位！")
        return
    
    # 获取卦名
    base_gua_name = get_gua_name_by_yao(gua_yao_mapping, hexagram)
    
    if not base_gua_name:
        print("无法从卦爻映射中找到卦名！")
        return

    # 提取卦辞内容
    base_gua_content = extract_guaci_content(docx_path, base_gua_name)
    if not base_gua_content:
        print(f"未找到名为 '{base_gua_name}' 的卦辞内容！")
        return

    print(f"本卦：{base_gua_name}")
    print(f"本卦卦辞：{base_gua_content}")
    
    # 计算变爻位置并生成变卦
    changing_line = calculate_changing_line(nums)  # 将 nums 传递给 calculate_changing_line
    changing_hexagram = apply_changing_line_to_hexagram(hexagram, changing_line)
    
    # 获取变卦的卦名
    changing_gua_name = get_gua_name_by_yao(gua_yao_mapping, changing_hexagram)
    changing_gua_content = extract_guaci_content(docx_path, changing_gua_name)
    if not changing_gua_content:
        print(f"未找到名为 '{changing_gua_name}' 的变卦卦辞！")
        return

    print(f"变卦：{changing_gua_name}")
    print(f"变卦卦辞：{changing_gua_content}")
    
    return {
        "base_gua": base_gua_name,
        "base_gua_content": base_gua_content,
        "changing_gua": changing_gua_name,
        "changing_gua_content": changing_gua_content,
    }
    
def main():
    # 指定卦爻的docx文件路径
    gua_yao_file_path = r"E:\code\forcast\易经自动算命\卦爻.docx"  # 替换为实际的文件路径
    # 指定《易经卦辞.docx》的路径
    docx_path = r"E:\code\forcast\易经自动算命\易经卦辞.docx"

    # 读取卦爻信息
    gua_yao_mapping = read_gua_yao_from_docx(gua_yao_file_path)


    # 询问用户想要的功能
    choice = input("请选择功能：\n1-生成每日一卦\n2-所求应物占卦占卜\n3-梅花数字占\n")
 
    if choice == "1":
        # 生成每日一卦
        daily_gua = generate_daily_hexagram_lunar(docx_path, gua_yao_mapping)
        print("今日本卦：", daily_gua["base_gua"])
        print("本卦卦辞：", daily_gua["base_gua_content"])
        print("变卦：", daily_gua["changing_gua"])
        print("变卦卦辞：", daily_gua["changing_gua_content"])  
        
        with open("shared_gua_names.txt", "w", encoding="utf-8") as f:
            base_gua_name_simplified = daily_gua["base_gua"].rstrip('卦')
            changing_gua_name_simplified = daily_gua["changing_gua"].rstrip('卦')
            f.write(f"{base_gua_name_simplified}\n{changing_gua_name_simplified}\n")

    elif choice == "2":
        # 询问用户想要了解的内容
        user_input = input("您想要问什么？")

        # 生成本卦
        base_hexagram = generate_base_hexagram_new()
        base_gua_name = get_gua_name_by_yao(gua_yao_mapping, base_hexagram)
        print("本卦：", base_hexagram)
        print("本卦卦名：", base_gua_name)

        # 提取本卦卦辞内容
        base_gua_content = extract_guaci_content(docx_path, base_gua_name)
        if base_gua_content:
            print(f"本卦卦辞：{base_gua_content}")
            print("--------------------")
        else:
            print(f"未找到名为'{base_gua_name}'的卦辞")

        # 生成变爻位
        changing_line, changing_line_index = generate_changing_line_new(base_hexagram)
        print("变爻位：第", changing_line_index+1, "爻")

        # 生成变卦
        changing_hexagram = generate_changing_hexagram(base_hexagram, changing_line, changing_line_index)
        changing_gua_name = get_gua_name_by_yao(gua_yao_mapping, changing_hexagram)
        print("变卦：", changing_hexagram)
        print("变卦卦名：", changing_gua_name)

        # 提取变卦卦辞内容
        changing_gua_content = extract_guaci_content(docx_path, changing_gua_name)
        if changing_gua_content:
            print(f"变卦卦辞：{changing_gua_content}")
            print("--------------------")
        else:
            print(f"未找到名为'{changing_gua_name}'的卦辞")

        # 写入卦名到文本文件
        if base_gua_name and changing_gua_name:
            with open("shared_gua_names.txt", "w", encoding="utf-8") as f:
                base_gua_name_simplified = base_gua_name.rstrip('卦')
                changing_gua_name_simplified = changing_gua_name.rstrip('卦')
                f.write(f"{base_gua_name_simplified}\n{changing_gua_name_simplified}\n")

            # Writing to the signature text file
            with open("签文.txt", "w", encoding="utf-8") as f:
                f.write(f"本卦卦名：{base_gua_name}\n")
                f.write(f"本卦卦辞：{base_gua_content}\n")
                f.write(f"变卦卦名：{changing_gua_name}\n")
                f.write(f"变卦卦辞：{changing_gua_content}\n")
        else:
            print("错误：无效的卦名，未写入文件。")

    elif choice == "3":
        # 梅花数字占
        base_gua_info = number_to_hexagram_from_user_input(docx_path, gua_yao_mapping) # 获取六爻卦信息
        if base_gua_info:
            print(f"本卦卦名：{base_gua_info['base_gua']}")
            print(f"本卦卦辞：{base_gua_info['base_gua_content']}") 
            print("--------------------")
            print(f"变卦卦名：{base_gua_info['changing_gua']}")
            print(f"变卦卦辞：{base_gua_info['changing_gua_content']}")
            # Write base and changing hexagram names to a file only if valid
            with open("shared_gua_names.txt", "w", encoding="utf-8") as f:
                base_gua_name_simplified = base_gua_info.get("base_gua", "").rstrip('卦')
                changing_gua_name_simplified = base_gua_info.get("changing_gua", "").rstrip('卦')
                if base_gua_name_simplified and changing_gua_name_simplified:
                    f.write(f"{base_gua_name_simplified}\n{changing_gua_name_simplified}\n")
                else:
                    print("Error: Missing valid hexagram names.")
        
                # Writing to the signature text file
                with open("签文.txt", "w", encoding="utf-8") as f:
                    f.write(f"本卦卦名：{base_gua_info['base_gua']}\n")
                    f.write(f"本卦卦辞：{base_gua_info['base_gua_content']}\n")
                    f.write(f"变卦卦名：{base_gua_info['changing_gua']}\n")
                    f.write(f"变卦卦辞：{base_gua_info['changing_gua_content']}\n")
    else:
        print("未生成卦名")
        


if __name__ == '__main__':
    main()
           
    # 询问用户是否需要了解白话文解释
    need_explanation = input("需要了解白话文解释吗？(y/n)：").strip().lower()
    if need_explanation == 'y':
        second_main()
